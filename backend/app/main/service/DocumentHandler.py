import re
from datetime import datetime

import pandas as pd

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
import tabula
import app.main.service.pdf_redactor as pdf_redactor

import docx
from docx.text.paragraph import Paragraph
from docx.table import Table

from bs4 import BeautifulSoup
from bs4.formatter import HTMLFormatter

from app.main.util.fileUtils import encode, itemIterator, markInHtml
from app.main.service.NameSearchByEntities import NameSearchByEntities
from app.main.util.heuristicMeasures import MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS, MEASURE_TO_COLUMN_KEY_REFERS_TO_NAMES
from app.main.service.languageBuilder import LanguageBuilder
from app.main.util.semanticWordLists import listOfVectorWords
from app.main.service.NamePickerInTables import NamePickerInTables


class DocumentHandler:

    def __init__(self, path: str, destiny: str = ""):
        self.path = path
        self.destiny = destiny
        self.nameSearch = NameSearchByEntities()

    def documentsProcessing(self):
        pass

    # def documentTagger(self):
    # pass

    def createFileOfName(self):
        self.createCsv(self.giveListNames())

    def giveListNames(self) -> list:
        pass

    def createCsv(self, listNames: list):
        dataFrame = pd.DataFrame(listNames, columns=['Names'])
        export_csv = dataFrame.to_csv(self.destiny, index=None, header=True)
        print(export_csv)


# TODO? optimize
class DocumentHandlerPdf(DocumentHandler):

    def __init__(self, path: str, destiny: str = ""):
        super().__init__(path, destiny=destiny)
        self.options = pdf_redactor.RedactorOptions()
        self.options.metadata_filters = {
            "Title": [lambda value: value],

            "Producer": [lambda value: value],
            "CreationDate": [lambda value: datetime.utcnow()],

            "DEFAULT": [lambda value: None],
        }
        self.options.xmp_filters = [lambda xml: None]

    def getPossibleColumnsNames(self, df):
        for key, typeColumn in zip(df.keys(), df.dtypes):
            if typeColumn == object:
                listOfWordSemantics = list(
                    filter(
                        lambda x: LanguageBuilder().semanticSimilarity(key, x) > MEASURE_TO_COLUMN_KEY_REFERS_TO_NAMES,
                        listOfVectorWords))
                if listOfWordSemantics:
                    yield key

    def getPersonalDataInTables(self, listNames:list):
        lastKeys = []
        for table in tabula.read_pdf(self.path, pages="all", multiple_tables=True):
            #table = table.loc[:, ~table.columns.str.contains('^Unnamed')]
            hasNewKeys = False
            for key in self.getPossibleColumnsNames(table):
                if(not hasNewKeys): lastKeys.clear()
                hasNewKeys = True
                dfNotNull = table[key][table[key].notnull()]
                countOfName = sum(list(map(lambda x: self.nameSearch.checkNameInDB(str(x)), dfNotNull)))
                if countOfName / len(dfNotNull) > MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS:
                    listNames[len(listNames):] = dfNotNull
                    lastKeys.append(list(table.keys()).index(key))
            if lastKeys and not hasNewKeys and list(table.keys()):
                for indexKey in lastKeys:
                    try:
                        dataKey = list(table.keys())[indexKey]
                    except IndexError as identifier:
                        print(identifier)
                        continue
                    dfNotNull = table[dataKey][table[dataKey].notnull()]
                    countOfName = sum(list(map(lambda x: self.nameSearch.checkNameInDB(str(x)), dfNotNull)))
                    if countOfName / len(dfNotNull) > MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS:
                        listNames.append(dataKey)
                        listNames[len(listNames):] = dfNotNull

    def getPersonalDataInTexts(self, listNames: list):
        fp = open(self.path, 'rb')
        parser = PDFParser(fp)
        fp.close()
        doc = PDFDocument()
        parser.set_document(doc)
        doc.set_parser(parser)
        doc.initialize('')
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        laparams.char_margin = 1.0
        laparams.word_margin = 1.0
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in doc.get_pages():
            interpreter.process_page(page)
            layout = device.get_result()
            for lt_obj in layout:
                if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                    for text in lt_obj.get_text().split("\n"):
                        if text and LanguageBuilder().hasContex(text):
                            doc = self.nameSearch.searchNames(text)
                            for entity in doc:
                                listNames.append(entity['name'].strip("\n"))

    def giveListNames(self) -> list:
        listNames = []
        self.getPersonalDataInTables(listNames)
        self.getPersonalDataInTexts(listNames)
        return listNames

    def documentsProcessing(self):
        listNames = list(set(self.giveListNames()))
        if not listNames:
            pdf_redactor.redactor(self.options, self.path, self.destiny)
            return
        listNames.sort(
                key=lambda value: len(value),
                reverse=True
            )
        regex = '|'.join(listNames)
        pdf_redactor.redactor(self.options, self.path, self.destiny)
        self.options.content_filters = [
            (
                re.compile(regex),
                lambda m: encode(m.group())
            )
        ]
        pdf_redactor.redactor(self.options, self.path, self.destiny)



class DocumentHandlerDocx(DocumentHandler):

    def __init__(self, path: str, destiny: str = ""):
        super().__init__(path, destiny=destiny)
        self.document = docx.Document(self.path)

    def getPickerData(self, table: Table) -> NamePickerInTables:
        picker = NamePickerInTables()
        isLabels = True
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if isLabels:
                        labels = list(
                            filter(lambda x: LanguageBuilder().semanticSimilarity(str(paragraph.text),
                                                                                  x) > MEASURE_TO_COLUMN_KEY_REFERS_TO_NAMES,
                                   listOfVectorWords))
                        if labels:
                            picker.addIndexColumn(index)
                    else:
                        if picker.isColumnName(index) and bool(paragraph.text.strip()):
                            picker.addName(index, paragraph.text)
                            if self.nameSearch.checkNameInDB(paragraph.text): picker.countRealName(index)
            isLabels = False
        return picker

    def definePicker(self, table: Table, picker: NamePickerInTables):
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                if picker.isColumnName(index):
                    for paragraph in cell.paragraphs:
                        if bool(paragraph.text.strip()):
                            picker.addName(index, paragraph.text)
                            if self.nameSearch.checkNameInDB(paragraph.text): picker.countRealName(index)

    def documentsProcessing(self):
        LastIndexesColumn = []
        for block in itemIterator(self.document):
            if isinstance(block, Paragraph):
                listNames = self.nameSearch.searchNames(block.text)
                for name in listNames:
                    regexName = re.compile(name['name'])
                    text = regexName.sub(encode(name['name']), block.text)
                    block.text = text
            elif isinstance(block, Table):
                picker = self.getPickerData(block)
                initialRow = 1
                if picker.getIndexesColumn():
                    LastIndexesColumn = picker.getIndexesColumn()
                elif LastIndexesColumn:
                    picker.addIndexesColumn(LastIndexesColumn)
                    self.definePicker(block, picker)
                    initialRow = 0
                else:
                    continue
                for row in block.rows[initialRow:]:
                    for index, cell in enumerate(row.cells):
                        if picker.isRealColumName(index, MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS):
                            for paragraph in cell.paragraphs:
                                paragraph.text = encode(paragraph.text)
            else:
                continue
        self.document.save(self.destiny)

    def giveListNames(self) -> list:
        LastIndexesColumn = []
        listNames = []
        for block in itemIterator(self.document):
            if isinstance(block, Paragraph):
                listOfMarks = self.nameSearch.searchNames(block.text)
                if listOfMarks != []:
                    listNames[len(listNames):] = [name['name'] for name in listOfMarks]
            elif isinstance(block, Table):
                picker = self.getPickerData(block)
                if picker.getIndexesColumn():
                    LastIndexesColumn = picker.getIndexesColumn()
                elif LastIndexesColumn:
                    picker.addIndexesColumn(LastIndexesColumn)
                    self.definePicker(block, picker)
                listNames[len(listNames):] = picker.getAllNames(MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS)

            else:
                continue
        return listNames


class DocumentHandlerExcel(DocumentHandler):

    def __init__(self, path: str, destiny: str = ""):
        super().__init__(path, destiny=destiny)
        self.df = pd.read_excel(path)

    def getPossibleColumnsNames(self):
        for key, typeColumn in zip(self.df.keys(), self.df.dtypes):
            if typeColumn == object:
                listOfWordSemantics = list(
                    filter(
                        lambda x: LanguageBuilder().semanticSimilarity(key, x) > MEASURE_TO_COLUMN_KEY_REFERS_TO_NAMES,
                        listOfVectorWords))
                if listOfWordSemantics:
                    yield key

    def documentsProcessing(self):
        for key in self.getPossibleColumnsNames():
            # print(key)
            dfNotNull = self.df[key][self.df[key].notnull()]
            countOfName = sum(list(map(lambda x: self.nameSearch.checkNameInDB(str(x)), dfNotNull)))
            if countOfName / len(dfNotNull) > MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS:
                self.df[key].replace({str(name): encode(str(name)) for name in dfNotNull}, inplace=True)
        self.df.to_excel(self.destiny, index=False)

    def giveListNames(self) -> list:
        listNames = []
        for key in self.getPossibleColumnsNames():
            # print(key)
            dfNotNull = self.df[key][self.df[key].notnull()]
            countOfName = sum(list(map(lambda x: self.nameSearch.checkNameInDB(str(x)), dfNotNull)))
            if countOfName / len(dfNotNull) > MEASURE_FOR_TEXTS_WITHOUT_CONTEXTS:
                listNames[len(listNames):] = dfNotNull
        return listNames


class DocumentHandlerHtml(DocumentHandler):

    def __init__(self, path: str, destiny: str = ""):
        super().__init__(path, destiny=destiny)
        with open(self.path, "r", encoding="utf8") as f:
            self.soup = BeautifulSoup(f.read(), "lxml")

    def locateNames(self, sentence):
        listNames = self.nameSearch.searchNames(str(sentence))
        if listNames is []:
            return sentence
        newSentence = ""
        index = 0
        for name in listNames:
            newSentence += sentence[index:name['star_char']] + markInHtml(name['name'])
            index = name['end_char']
        if index <= len(sentence) - 1:
            newSentence += sentence[index:]
        return newSentence

    def encodeNames(self, sentence):
        print(sentence)
        listNames = self.nameSearch.searchNames(str(sentence))
        if listNames is []:
            return sentence
        newSentence = ""
        index = 0
        for name in listNames:
            newSentence += sentence[index:name['star_char']] + encode(name['name'])
            index = name['end_char']
        if index <= len(sentence) - 1:
            newSentence += sentence[index:]
        return newSentence

    def documentsProcessing(self):
        formatter = HTMLFormatter(self.encodeNames)
        with open(self.destiny, "w") as f:
            f.write(self.soup.prettify(formatter=formatter))

    def documentTagger(self):
        formatter = HTMLFormatter(self.locateNames)
        with open(self.destiny, "w") as f:
            f.write(self.soup.prettify(formatter=formatter))

    def giveListNames(self):
        listNames = []
        blacklist = ['[document]', 'noscript', 'header',
                     'html', 'meta', 'head', 'input', 'script', 'link', 'lang']
        for lable in self.soup.stripped_strings:
            if lable not in blacklist:
                # print(lable)
                listOfMarks = self.nameSearch.searchNames(str(lable))
                listNames[len(listNames):] = [name['name'].replace("\n", "") for name in listOfMarks]
        return listNames


class DocumentHandlerTxt(DocumentHandler):

    def modifyLine(self, line: str, listNames: list) -> str:
        newLine = ""
        index = 0
        for name in listNames:
            newLine += line[index:name['star_char']] + encode(name['name'])
            index = name['end_char']
        if index <= len(line) - 1:
            newLine += line[index:]
        return newLine

    def documentsProcessing(self):
        with open(self.path, 'r') as file, open(self.destiny, 'w') as destiny:
            for line in file:
                listNames = self.nameSearch.searchNames(line)
                destiny.writelines(self.modifyLine(line, listNames))

    def giveListNames(self) -> list:
        listNames = []
        with open(self.path, 'r') as file:
            for line in file:
                listNames[len(listNames):] = [name['name'] for name in self.nameSearch.searchNames(line)]
        return listNames
