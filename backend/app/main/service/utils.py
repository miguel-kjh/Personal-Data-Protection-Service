from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
import random
from typing import Text



def proc_pdf3k(document):
    fp = open(document, 'rb')
    parser = PDFParser(fp)
    fp.close()
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    doc.initialize('')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    laparams.char_margin = 1.0
    laparams.word_margin = 1.0
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    extracted_text = ''

    for page in doc.get_pages():
        interpreter.process_page(page)
        layout = device.get_result()
        for lt_obj in layout:
            if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                print(lt_obj.get_text())
                print("----------")
                extracted_text += lt_obj.get_text()


def encode(text:str):
    return "*"*len(text)

def markInHtml(text:str):
    return '<mark style="background: #7aecec;">' + text + '</mark>'


def run_append(paragraph, r, text, bold=False):
    run = paragraph.add_run(text, r.style)
    if bold:
        run.bold = True
    else:
        run.bold = r.bold
    run.font.name = r.font.name
    run.font.size = r.font.size
    run.italic = r.italic
    run.underline = r.underline


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def proc_docx(document):
    doc = docx.Document(document)
    for block in iter_block_items(doc):
        if block is None:
            continue
        if isinstance(block, Paragraph):
            print(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                row_data = []
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        row_data.append(paragraph.text)
                print("\t".join(row_data))

