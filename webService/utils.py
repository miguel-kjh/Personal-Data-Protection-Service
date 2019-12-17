from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from app import ALLOWED_EXTENSIONS
import random
from datetime import datetime
import hashlib
from unidecode import unidecode
from typing import Text



def proc_pdf3k(document):
    fp = open(document, 'rb')
    parser = PDFParser(fp)
    fp.close()
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    doc.initialize('')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    laparams.char_margin = 1.0
    laparams.word_margin = 1.0
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    extracted_text = ''

    for page in doc.get_pages():
        interpreter.process_page(page)
        layout = device.get_result()
        for lt_obj in layout:
            if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                print(lt_obj.get_text())
                print("----------")
                extracted_text += lt_obj.get_text()


def encode(text:str):
    return "*"*len(text)

def markInHtml(text:str):
    return '<mark style="background: #7aecec;">' + text + '</mark>'


def run_append(paragraph, r, text, bold=False):
    run = paragraph.add_run(text, r.style)
    if bold:
        run.bold = True
    else:
        run.bold = r.bold
    run.font.name = r.font.name
    run.font.size = r.font.size
    run.italic = r.italic
    run.underline = r.underline


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def proc_docx(document):
    doc = docx.Document(document)
    for block in iter_block_items(doc):
        if block is None:
            continue
        if isinstance(block, Paragraph):
            print(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                row_data = []
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        row_data.append(paragraph.text)
                print("\t".join(row_data))

def allowedFile(filename:str) -> bool:
	return giveTypeOfFile(filename) in ALLOWED_EXTENSIONS

def giveTypeOfFile(filename:str) -> str:
    return '.' in filename and filename.rsplit('.', 1)[1].lower()

def giveFileNameUnique(filename:str, fileType:str) -> str:
    #filename = filename + str(datetime.now().timestamp())
    #sha = hashlib.sha256(filename.encode())
    #return sha.hexdigest() + "." + fileType
    return str(datetime.now().timestamp()).replace(".","") + "." + fileType

def normalizeUnicode(string:str) -> str: 
    return unidecode(string)

def cleanHeadAndTailOfList(listTokens:list):
    for token in reversed(listTokens):
        if token[1].pos_ == "PROPN": break
        listTokens.remove(token)
    for token in listTokens:
        if token[1].pos_ == "PROPN": break
        listTokens.remove(token)

def generatorNames(nlp, text:Text):
    with nlp.disable_pipes('parser','ner'):
        doc = nlp(text)
        articules = ["de", "del","-","el","los","todos"]
        listTokens = [(index,token) for index,token in enumerate(doc) if token.pos_ == 'PROPN' or token.text.lower() in articules]
        cleanHeadAndTailOfList(listTokens)
        if listTokens == []: return listTokens
        names = [listTokens[0]]
        if len(listTokens) == 1 and names[0][1].pos_ == 'PROPN':
            yield names
        else:
            countNames = 0
            for token in listTokens[1:]:
                if token[0] == names[countNames][0] + 1:
                    names.append(token)
                    if listTokens[-1] == token:
                        if names[0][1].pos_ == 'PROPN' and names[-1][1].pos_ == names[0][1].pos_:
                            yield names
                        break    
                    countNames += 1
                else:
                    if names[0][1].pos_ == 'PROPN' and names[-1][1].pos_ == names[0][1].pos_:
                        yield names
                    names = []
                    names.append(token)
                    countNames = 0
