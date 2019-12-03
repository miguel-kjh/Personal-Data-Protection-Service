import re
from datetime import datetime

import pandas as pd

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
import pdf_redactor

import docx
from docx.text.paragraph import Paragraph
from docx.table import Table

from bs4 import BeautifulSoup
from bs4.formatter import HTMLFormatter

from utils import proc_pdf3k, proc_docx, run_append, encode, iter_block_items, markInHtml
from SearcherNamesTexts import SearcherNamesTexts


class DocumentHandler():

    def __init__(self, path:str, nlp, destiny:str = ""):
        self.document = path
        self.destiny = destiny
        self.searcherNamesTexts = SearcherNamesTexts(nlp)

    def read(self):
        pass

    def documentsProcessing(self):
        pass

    #def documentTagger(self):
        #pass

    def createFileOfName(self):
        self.createCsv(self.giveListNames())
        return self.destiny

    def giveListNames(self):
        pass

    def createCsv(self, listNames:list):
        dataFrame = pd.DataFrame(listNames, columns=['Names'])
        export_csv = dataFrame.to_csv(self.destiny, index = None, header=True)
        print(export_csv)

#TODO? optimize
class DocumentHandlerPDF(DocumentHandler):

    def __init__(self, path:str, nlp,destiny:str = ""):
        super().__init__(path,nlp,destiny=destiny)
        self.options = pdf_redactor.RedactorOptions()
        self.options.metadata_filters = {
            "Title": [lambda value: value],

            "Producer": [lambda value: value],
            "CreationDate": [lambda value: datetime.utcnow()],

            "DEFAULT": [lambda value: None],
        }
        self.options.xmp_filters = [lambda xml: None]

    def read(self):
        proc_pdf3k(self.document)

    def giveListNames(self) -> list:
        fp = open(self.document, 'rb')
        parser = PDFParser(fp)
        fp.close()
        doc = PDFDocument()
        parser.set_document(doc)
        doc.set_parser(parser)
        doc.initialize('')
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        laparams.char_margin = 1.0
        laparams.word_margin = 1.0
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        listNames = []
        for page in doc.get_pages():
            interpreter.process_page(page)
            layout = device.get_result()
            for lt_obj in layout:
                if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                    doc = self.searcherNamesTexts.searchNames(lt_obj.get_text())
                    for e in doc:
                        listNames.append(e['name'].strip("\n"))
        return list(set(listNames))

    def documentsProcessing(self):
        listNames = self.giveListNames()
        print(listNames[:])
        if listNames is not []:
            listNames.sort(
                key=lambda value: len(value),
                reverse=True
            )
            self.options.content_filters = [
                (
                    re.compile(listNames[0]),
                    lambda m: encode(listNames[0]).upper()
                )
            ]
            pdf_redactor.redactor(self.options, self.document, self.destiny)
            for name in listNames[1:]:
                self.options.content_filters = [
                    (
                        re.compile(name),
                        lambda m: encode(name).upper()
                    )
                ]
                pdf_redactor.redactor(self.options, self.destiny, self.destiny)
        


class DocumentHandlerDocx(DocumentHandler):

    def read(self):
        return proc_docx(self.document)

    #TODO? extend this to other docx's objects like images
    def documentsProcessing(self):
        doc = docx.Document(self.document)
        document = docx.Document()
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                paragraph = document.add_paragraph()
                self.localizeNames(block,paragraph)
            elif isinstance(block, Table):
                table = document.add_table(rows=len(block.rows), cols=len(block.columns))
                table.style = block.style
                for index_row,row in enumerate(block.rows):
                    for index_cell,cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            paragraph_table = table.cell(index_row,index_cell).add_paragraph()
                            self.localizeNames(paragraph,paragraph_table)
            else:
                continue
        document.save(self.destiny)

    def localizeNames(self,block,paragraph):
        for r in block.runs:
            listOfMarks = self.searcherNamesTexts.searchNames(r.text)
            if listOfMarks not in []:
                index = 0
                for ele in listOfMarks:
                    run_append(paragraph, r, r.text[index:ele['star_char']])
                    run_append(paragraph, r, encode(r.text[ele['star_char']:ele['end_char']]), True)
                    index = ele['end_char']
                if index <= len(r.text) - 1:
                    run_append(paragraph, r, r.text[index:])
            else:
                run_append(paragraph, r, r.text)

    def giveListNames(self):
        doc = docx.Document(self.document)
        document = docx.Document()
        listNames = []
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                listOfMarks = self.searcherNamesTexts.searchNames(block.text)
                if listOfMarks != []:
                    listNames[len(listNames):] = [name['name'] for name in listOfMarks]
            elif isinstance(block, Table):
                table = document.add_table(rows=len(block.rows), cols=len(block.columns))
                table.style = block.style
                for row in block.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            listOfMarks = self.searcherNamesTexts.searchNames(paragraph.text)
                            if listOfMarks != []:
                                listNames[len(listNames):] = [name['name'] for name in listOfMarks]
            else:
                continue
        return listNames
        

class DocumentHandlerExe(DocumentHandler):

    def __init__(self,path:str,nlp,destiny:str = ""):
        super().__init__(path,nlp,destiny=destiny)
        self.df = pd.read_excel(path)

    def read(self):
        print(self.df)

    #TODO? save formules.
    def documentsProcessing(self):
        for key in self.df.keys():
            for index,ele in enumerate(self.df[key]):
                if self.searcherNamesTexts.isName(str(ele)):
                    self.df.at[index,key] = encode(self.df.at[index,key])
        self.df.to_excel(self.destiny)

    def giveListNames(self):
        listNames = []
        for key in self.df.keys():
            for ele in self.df[key]:
                if self.searcherNamesTexts.isName(str(ele)):
                    listNames.append(ele)
        return listNames

class DocumentHandlerHTML(DocumentHandler):

    def __init__(self,path:str,nlp,destiny:str = ""):
        super().__init__(path,nlp,destiny=destiny)
        with open(self.document,"r") as f:
            self.soup = BeautifulSoup(f.read(), "lxml")

    def read(self):
        print(self.soup)

    def locateNames(self,sentence):
        listNames = self.searcherNamesTexts.searchNames(str(sentence))
        if listNames is []:
            return sentence
        newSentence = ""
        index = 0
        for name in listNames:
            newSentence += sentence[index:name['star_char']] + markInHtml(name['name'])
            index = name['end_char']
        if index <= len(sentence)-1:
            newSentence += sentence[index:]
        return newSentence

    def encodeNames(self,sentence):
        listNames = self.searcherNamesTexts.searchNames(str(sentence))
        if listNames is []:
            return sentence
        newSentence = ""
        index = 0
        for name in listNames:
            newSentence += sentence[index:name['star_char']] + encode(name['name'])
            index = name['end_char']
        if index <= len(sentence)-1:
            newSentence += sentence[index:]
        return newSentence

    def documentsProcessing(self):
        formatter = HTMLFormatter(self.encodeNames)
        with open(self.destiny,"w") as f:
            f.write(self.soup.prettify(formatter=formatter))

    def documentTagger(self):
        formatter = HTMLFormatter(self.locateNames)
        with open(self.destiny,"w") as f:
            f.write(self.soup.prettify(formatter=formatter))

    def giveListNames(self):
        listNames = []
        blacklist = [ '[document]' , 'noscript' , 'header' , 
        'html' , 'meta' , 'head' , 'input' , 'script', 'link'] # convertir esto en regex
        for lable in self.soup.find_all(text=True):
            if lable not in blacklist:
                #print(lable)
                listOfMarks = self.searcherNamesTexts.searchNames(str(lable))
                listNames[len(listNames):] = [name['name'].replace("\n","") for name in listOfMarks]
        return listNames



